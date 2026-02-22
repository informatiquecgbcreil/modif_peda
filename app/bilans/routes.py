from app.rbac import require_perm
import datetime
import json
import os
import uuid
from io import BytesIO

from flask import Blueprint, render_template, request, abort, redirect, url_for, flash, send_file, current_app
from flask_login import login_required, current_user
from docx import Document
from docx.shared import Inches
from werkzeug.utils import secure_filename

from app.bilans.services import (
    compute_alertes,
    compute_depenses_mensuelles,
    compute_depenses_par_secteur,
    compute_kpis,
    compute_bilan_secteur,
    compute_bilan_subvention,
    compute_qualite_gestion,
    compute_stats_inventaire,
    compute_bilans_lourds,
    list_secteurs,
    list_subventions,
    list_exercice_years,
    scope_for_user,
)
from app.extensions import db
from app.models import BilanLourdNarratif
from app.services.storage import ensure_upload_subdir, media_relpath


bp = Blueprint("bilans", __name__, url_prefix="")


@bp.route("/bilans")
@login_required
@require_perm("bilans:view")
def dashboard():
    scope = scope_for_user(current_user)

    years = list_exercice_years(scope)
    # année sélectionnée : param ?year=YYYY, sinon la plus récente dispo
    year_param = request.args.get("year")
    try:
        year = int(year_param) if year_param else years[0]
    except (TypeError, ValueError):
        year = years[0]
    if year not in years:
        # évite de forcer un year arbitraire via l'URL
        abort(403)

    kpis = compute_kpis(year, scope)
    series = compute_depenses_mensuelles(year, scope)
    par_secteur = compute_depenses_par_secteur(year, scope)
    alertes = compute_alertes(year, scope)

    multi_secteurs = scope.secteurs is None

    return render_template(
        "bilans_dashboard.html",
        year=year,
        years=years,
        kpis=kpis,
        series=series,
        par_secteur=par_secteur,
        alertes=alertes,
        multi_secteurs=multi_secteurs,
        scope=scope,
    )




def _notes_secteur_key(scope):
    if scope.secteurs is None:
        return "__ALL__"
    if scope.secteurs:
        return scope.secteurs[0]
    return "__NONE__"


def _get_or_create_bilan_note(year: int, scope, create: bool = False):
    secteur_key = _notes_secteur_key(scope)
    note = BilanLourdNarratif.query.filter_by(annee=year, secteur=secteur_key).first()
    if not note and create:
        note = BilanLourdNarratif(annee=year, secteur=secteur_key)
        db.session.add(note)
    return note, secteur_key


def _safe_json_list(raw_value: str | None):
    try:
        data = json.loads(raw_value or "[]")
        return data if isinstance(data, list) else []
    except Exception:
        return []


def _parse_timeline(raw_text: str) -> list[dict]:
    rows = []
    for line in (raw_text or "").splitlines():
        line = line.strip()
        if not line:
            continue
        parts = [p.strip() for p in line.split("|", 2)]
        if len(parts) == 1:
            rows.append({"date": "", "titre": parts[0], "detail": ""})
        elif len(parts) == 2:
            rows.append({"date": parts[0], "titre": parts[1], "detail": ""})
        else:
            rows.append({"date": parts[0], "titre": parts[1], "detail": parts[2]})
    return rows


def _timeline_text(rows: list[dict]) -> str:
    out = []
    for row in rows:
        date_label = row.get("date") or ""
        title = row.get("titre") or ""
        detail = row.get("detail") or ""
        out.append(f"{date_label} | {title} | {detail}".strip(" |"))
    return "\n".join(out)


def _save_bilan_photos(files, year: int, scope_key: str) -> list[dict]:
    uploads = []
    allowed = {".jpg", ".jpeg", ".png", ".webp", ".gif"}
    safe_scope = secure_filename(scope_key or "scope") or "scope"
    target_dir = ensure_upload_subdir("bilans_lourds", str(year), safe_scope)

    for f in files:
        if not f or not f.filename:
            continue
        original = secure_filename(f.filename)
        _, ext = os.path.splitext(original.lower())
        if ext not in allowed:
            continue
        out_name = f"{uuid.uuid4().hex[:10]}_{original}"
        out_abs = os.path.join(target_dir, out_name)
        f.save(out_abs)
        rel = media_relpath("bilans_lourds", str(year), safe_scope, out_name)
        uploads.append({"path": rel, "caption": ""})
    return uploads


@bp.route("/bilans/lourds", methods=["GET", "POST"])
@login_required
@require_perm("bilans:view")
def bilans_lourds():
    scope = scope_for_user(current_user)
    years = list_exercice_years(scope) or [datetime.date.today().year]

    year_param = request.args.get("year")
    try:
        year = int(year_param) if year_param else years[0]
    except (TypeError, ValueError, IndexError):
        year = years[0]

    # Au lieu d'un 403, on redirige vers une année valide
    if year not in years:
        return redirect(url_for("bilans.bilans_lourds", year=years[0]))

    if request.method == "POST":
        action = request.form.get("action")
        if action == "save_notes":
            note, _ = _get_or_create_bilan_note(year, scope, create=True)
            note.faits_marquants = (request.form.get("faits_marquants") or "").strip() or None
            note.difficultes = (request.form.get("difficultes") or "").strip() or None
            note.perspectives = (request.form.get("perspectives") or "").strip() or None

            timeline_rows = _parse_timeline(request.form.get("timeline_input") or "")
            note.timeline_json = json.dumps(timeline_rows, ensure_ascii=False) if timeline_rows else None

            current_photos = _safe_json_list(note.photos_json)
            remove_paths = {p for p in request.form.getlist("remove_photo") if p}
            kept = [p for p in current_photos if p.get("path") not in remove_paths]
            uploaded = _save_bilan_photos(request.files.getlist("photos"), year, _notes_secteur_key(scope))
            final_photos = kept + uploaded
            note.photos_json = json.dumps(final_photos, ensure_ascii=False) if final_photos else None

            note.updated_by_user_id = current_user.id
            if not note.created_by_user_id:
                note.created_by_user_id = current_user.id
            db.session.commit()
            flash("Champs libres enregistrés.", "success")
        return redirect(url_for("bilans.bilans_lourds", year=year))

    stats = compute_bilans_lourds(year, scope)
    note, note_scope = _get_or_create_bilan_note(year, scope, create=False)
    photos = _safe_json_list(note.photos_json if note else None)
    timeline_rows = _safe_json_list(note.timeline_json if note else None)
    prev = stats.get("comparatif", {}).get("activite", {})
    prev_eval = stats.get("comparatif", {}).get("evaluations", {})

    def _delta_pct(cur: float, old: float):
        if not old:
            return None
        return round(((cur - old) / old) * 100, 1)

    deltas = {
        "sessions": _delta_pct(stats["activite"].get("nb_sessions", 0), prev.get("nb_sessions", 0)),
        "presences": _delta_pct(stats["activite"].get("nb_presences", 0), prev.get("nb_presences", 0)),
        "participants_uniques": _delta_pct(stats["activite"].get("nb_participants_uniques", 0), prev.get("nb_participants_uniques", 0)),
        "evaluations": _delta_pct(stats["evaluations"].get("total", 0), prev_eval.get("total", 0)),
    }

    executive_summary = [
        f"{stats['activite'].get('nb_sessions', 0)} sessions réalisées pour {stats['activite'].get('nb_presences', 0)} présences.",
        f"{stats['activite'].get('nb_participants_uniques', 0)} participants uniques, dont {stats['activite'].get('nb_participants_retour', 0)} reviennent au moins 2 fois ({stats['activite'].get('taux_fidelisation', 0)}%).",
        f"Intensité moyenne d'accompagnement : {stats['activite'].get('intensite_accompagnement', 0)} séance(s) par participant.",
        f"Taux de remplissage collectif : {stats['activite'].get('taux_remplissage_collectif', 0)}% et {stats['activite'].get('nb_rdv', 0)} RDV individuels ({stats['activite'].get('minutes_rdv', 0)} min).",
        f"{stats['evaluations'].get('total', 0)} évaluations réalisées sur {stats['evaluations'].get('nb_competences_uniques', 0)} compétence(s) unique(s).",
    ]

    return render_template(
        "bilans_lourds.html",
        year=year,
        years=years,
        stats=stats,
        deltas=deltas,
        executive_summary=executive_summary,
        narrative=note,
        narrative_scope=note_scope,
        narrative_photos=photos,
        timeline_rows=timeline_rows,
        timeline_input=_timeline_text(timeline_rows),
    )


@bp.route("/bilans/lourds/export.docx")
@login_required
@require_perm("bilans:view")
def bilans_lourds_export_docx():
    scope = scope_for_user(current_user)
    years = list_exercice_years(scope) or [datetime.date.today().year]

    year_param = request.args.get("year")
    try:
        year = int(year_param) if year_param else years[0]
    except (TypeError, ValueError, IndexError):
        year = years[0]

    if year not in years:
        return redirect(url_for("bilans.bilans_lourds", year=years[0]))

    stats = compute_bilans_lourds(year, scope)
    note, note_scope = _get_or_create_bilan_note(year, scope, create=False)

    doc = Document()
    doc.add_heading(f"Bilans lourds — Rapport {year}", level=1)
    scope_label = "Tous secteurs" if note_scope == "__ALL__" else note_scope
    doc.add_paragraph(f"Périmètre : {scope_label}")

    doc.add_heading("Synthèse activité", level=2)
    activite = stats.get("activite", {})
    for label, value in [
        ("Ateliers actifs", activite.get("nb_ateliers", 0)),
        ("Sessions réalisées", activite.get("nb_sessions", 0)),
        ("Présences", activite.get("nb_presences", 0)),
        ("Participants uniques", activite.get("nb_participants_uniques", 0)),
        ("Participants fidèles (>=2 venues)", activite.get("nb_participants_retour", 0)),
        ("Taux de fidélisation", f"{activite.get('taux_fidelisation', 0)}%"),
        ("Intensité d'accompagnement", f"{activite.get('intensite_accompagnement', 0)} séance(s)/participant"),
        ("Taux de remplissage collectif", f"{activite.get('taux_remplissage_collectif', 0)}%"),
        ("RDV individuels", activite.get("nb_rdv", 0)),
    ]:
        doc.add_paragraph(f"- {label} : {value}")

    doc.add_heading("Résultats pédagogiques", level=2)
    evaluations = stats.get("evaluations", {})
    doc.add_paragraph(f"- Évaluations : {evaluations.get('total', 0)}")
    doc.add_paragraph(f"- Compétences évaluées : {evaluations.get('nb_competences_uniques', 0)}")
    for label, nb in evaluations.get("par_etat", []):
        doc.add_paragraph(f"  • {label} : {nb}")

    doc.add_heading("Narratif", level=2)
    doc.add_paragraph(f"Faits marquants : {(note.faits_marquants if note else '') or 'Non renseigné'}")
    doc.add_paragraph(f"Difficultés rencontrées : {(note.difficultes if note else '') or 'Non renseigné'}")
    doc.add_paragraph(f"Perspectives / actions : {(note.perspectives if note else '') or 'Non renseigné'}")

    timeline_rows = _safe_json_list(note.timeline_json if note else None)
    if timeline_rows:
        doc.add_heading("Frise chronologique", level=2)
        for row in timeline_rows:
            date_label = row.get("date") or "Date"
            title = row.get("titre") or "Événement"
            detail = row.get("detail") or ""
            doc.add_paragraph(f"- {date_label} — {title}{': ' + detail if detail else ''}")

    photos = _safe_json_list(note.photos_json if note else None)
    if photos:
        doc.add_heading("Photographies marquantes", level=2)
        for p in photos:
            rel_path = (p.get("path") or "").strip()
            if not rel_path:
                continue

            abs_path = os.path.join(current_app.static_folder, rel_path)
            if os.path.exists(abs_path):
                try:
                    doc.add_picture(abs_path, width=Inches(5.8))
                except Exception:
                    doc.add_paragraph(f"- Image non insérable : {rel_path}")
            else:
                doc.add_paragraph(f"- Image introuvable : {rel_path}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    filename = f"bilans_lourds_{year}.docx"
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@bp.route("/bilans/secteur")
@login_required
@require_perm("bilans:view")
def bilan_secteur():
    scope = scope_for_user(current_user)
    years = list_exercice_years(scope)
    year_param = request.args.get("year")
    try:
        year = int(year_param) if year_param else years[0]
    except (TypeError, ValueError):
        year = years[0]
    if year not in years:
        abort(403)

    secteurs = list_secteurs(year, scope)

    # Choix secteur (finance/direction) via query param; responsable_secteur = auto
    selected = request.args.get("secteur")
    if scope.secteurs is not None:
        # responsable_secteur
        selected = scope.secteurs[0] if scope.secteurs else None
    if selected and selected not in secteurs:
        abort(403)
    if not selected and secteurs:
        selected = secteurs[0]

    data = compute_bilan_secteur(year, selected, scope) if selected else None

    return render_template(
        "bilans_secteur.html",
        year=year,
        years=years,
        secteurs=secteurs,
        selected_secteur=selected,
        data=data,
        scope=scope,
    )


@bp.route("/bilans/subvention")
@login_required
@require_perm("bilans:view")
def bilan_subvention():
    scope = scope_for_user(current_user)
    years = list_exercice_years(scope)
    year_param = request.args.get("year")
    try:
        year = int(year_param) if year_param else years[0]
    except (TypeError, ValueError):
        year = years[0]
    if year not in years:
        abort(403)

    subventions = list_subventions(year, scope)

    # Choix subvention (id) via query param
    selected_id = request.args.get("id")
    selected = None
    if selected_id:
        try:
            sid = int(selected_id)
        except ValueError:
            sid = None
        if sid:
            selected = next((s for s in subventions if s["id"] == sid), None)
            if not selected:
                abort(403)
    if not selected and subventions:
        selected = subventions[0]

    data = compute_bilan_subvention(year, selected["id"], scope) if selected else None

    return render_template(
        "bilans_subvention.html",
        year=year,
        years=years,
        subventions=subventions,
        selected_subvention=selected,
        data=data,
        scope=scope,
    )


@bp.route("/bilans/qualite")
@login_required
@require_perm("bilans:view")
def qualite():
    scope = scope_for_user(current_user)
    years = list_exercice_years(scope)
    year_param = request.args.get("year")
    try:
        year = int(year_param) if year_param else years[0]
    except (TypeError, ValueError):
        year = years[0]
    if year not in years:
        abort(403)

    data = compute_qualite_gestion(year, scope)
    return render_template("bilans_qualite.html", year=year, years=years, data=data, scope=scope)


@bp.route("/bilans/inventaire")
@login_required
@require_perm("bilans:view")
def inventaire():
    scope = scope_for_user(current_user)
    years = list_exercice_years(scope)
    year_param = request.args.get("year")
    try:
        year = int(year_param) if year_param else years[0]
    except (TypeError, ValueError):
        year = years[0]
    if year not in years:
        abort(403)

    data = compute_stats_inventaire(year, scope)
    return render_template("bilans_inventaire.html", year=year, years=years, data=data, scope=scope)
